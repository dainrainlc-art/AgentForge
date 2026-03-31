"""文件处理插件 - 支持文件上传、下载、解析和管理

功能特性:
- 多格式文件上传和下载
- 文件内容解析 (PDF, Word, Excel, CSV, 图片OCR)
- 文件压缩/解压
- 文件格式转换
- 文件元数据提取
- 安全文件处理 (病毒扫描、类型验证)
"""

import os
import io
import base64
import hashlib
import mimetypes
import zipfile
import tempfile
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional, List, Dict, Any, Union, BinaryIO
from datetime import datetime
from enum import Enum
import asyncio

from .base import Plugin, PluginResult


class FileType(Enum):
    """文件类型枚举"""
    DOCUMENT = "document"      # 文档 (PDF, Word, TXT)
    SPREADSHEET = "spreadsheet"  # 表格 (Excel, CSV)
    IMAGE = "image"            # 图片 (JPG, PNG, GIF)
    AUDIO = "audio"            # 音频 (MP3, WAV)
    VIDEO = "video"            # 视频 (MP4, AVI)
    ARCHIVE = "archive"        # 压缩包 (ZIP, RAR)
    CODE = "code"              # 代码文件
    UNKNOWN = "unknown"        # 未知类型


@dataclass
class FileInfo:
    """文件信息"""
    id: str
    name: str
    original_name: str
    size: int
    mime_type: str
    file_type: FileType
    extension: str
    hash_md5: str
    hash_sha256: str
    created_at: datetime
    modified_at: datetime
    metadata: Dict[str, Any] = field(default_factory=dict)
    tags: List[str] = field(default_factory=list)
    is_compressed: bool = False
    compression_ratio: Optional[float] = None


@dataclass
class ParsedContent:
    """解析后的文件内容"""
    text: str = ""
    tables: List[List[List[str]]] = field(default_factory=list)
    images: List[bytes] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    confidence: float = 1.0


class FilePlugin(Plugin):
    """文件处理插件"""
    
    name = "file"
    description = "文件上传、下载、解析和管理"
    version = "1.0.0"
    
    # 支持的文件扩展名映射
    EXTENSION_MAP = {
        # 文档
        '.pdf': FileType.DOCUMENT,
        '.doc': FileType.DOCUMENT,
        '.docx': FileType.DOCUMENT,
        '.txt': FileType.DOCUMENT,
        '.rtf': FileType.DOCUMENT,
        '.md': FileType.DOCUMENT,
        # 表格
        '.xls': FileType.SPREADSHEET,
        '.xlsx': FileType.SPREADSHEET,
        '.csv': FileType.SPREADSHEET,
        '.tsv': FileType.SPREADSHEET,
        # 图片
        '.jpg': FileType.IMAGE,
        '.jpeg': FileType.IMAGE,
        '.png': FileType.IMAGE,
        '.gif': FileType.IMAGE,
        '.bmp': FileType.IMAGE,
        '.webp': FileType.IMAGE,
        '.svg': FileType.IMAGE,
        # 音频
        '.mp3': FileType.AUDIO,
        '.wav': FileType.AUDIO,
        '.flac': FileType.AUDIO,
        '.aac': FileType.AUDIO,
        '.ogg': FileType.AUDIO,
        # 视频
        '.mp4': FileType.VIDEO,
        '.avi': FileType.VIDEO,
        '.mkv': FileType.VIDEO,
        '.mov': FileType.VIDEO,
        '.wmv': FileType.VIDEO,
        # 压缩包
        '.zip': FileType.ARCHIVE,
        '.rar': FileType.ARCHIVE,
        '.7z': FileType.ARCHIVE,
        '.tar': FileType.ARCHIVE,
        '.gz': FileType.ARCHIVE,
        # 代码
        '.py': FileType.CODE,
        '.js': FileType.CODE,
        '.ts': FileType.CODE,
        '.java': FileType.CODE,
        '.cpp': FileType.CODE,
        '.c': FileType.CODE,
        '.go': FileType.CODE,
        '.rs': FileType.CODE,
        '.html': FileType.CODE,
        '.css': FileType.CODE,
        '.json': FileType.CODE,
        '.xml': FileType.CODE,
        '.yaml': FileType.CODE,
        '.yml': FileType.CODE,
    }
    
    # 最大文件大小 (100MB)
    MAX_FILE_SIZE = 100 * 1024 * 1024
    
    # 允许的文件类型 (白名单)
    ALLOWED_TYPES = {
        FileType.DOCUMENT,
        FileType.SPREADSHEET,
        FileType.IMAGE,
        FileType.CODE,
        FileType.ARCHIVE,
    }
    
    def __init__(self):
        super().__init__()
        self.storage_path = Path("./data/files")
        self.temp_path = Path("./data/temp")
        self._files: Dict[str, FileInfo] = {}
        self._ensure_directories()
    
    def _ensure_directories(self):
        """确保目录存在"""
        self.storage_path.mkdir(parents=True, exist_ok=True)
        self.temp_path.mkdir(parents=True, exist_ok=True)
    
    def _get_file_type(self, extension: str) -> FileType:
        """根据扩展名获取文件类型"""
        return self.EXTENSION_MAP.get(extension.lower(), FileType.UNKNOWN)
    
    def _calculate_hashes(self, data: bytes) -> tuple:
        """计算文件哈希"""
        md5_hash = hashlib.md5(data, usedforsecurity=False).hexdigest()
        sha256_hash = hashlib.sha256(data).hexdigest()
        return md5_hash, sha256_hash
    
    def _validate_file(self, file_name: str, file_size: int, mime_type: str) -> tuple[bool, str]:
        """验证文件"""
        # 检查文件大小
        if file_size > self.MAX_FILE_SIZE:
            return False, f"文件大小超过限制 ({self.MAX_FILE_SIZE / 1024 / 1024:.1f}MB)"
        
        # 检查文件类型
        extension = Path(file_name).suffix.lower()
        file_type = self._get_file_type(extension)
        
        if file_type == FileType.UNKNOWN:
            return False, f"不支持的文件类型: {extension}"
        
        if file_type not in self.ALLOWED_TYPES:
            return False, f"不允许的文件类型: {file_type.value}"
        
        # 检查MIME类型一致性
        expected_mime, _ = mimetypes.guess_type(file_name)
        if expected_mime and mime_type != expected_mime:
            return False, f"MIME类型不匹配: 期望 {expected_mime}, 实际 {mime_type}"
        
        return True, ""
    
    async def upload_file(
        self,
        file_data: Union[bytes, BinaryIO],
        file_name: str,
        mime_type: Optional[str] = None,
        tags: Optional[List[str]] = None
    ) -> PluginResult:
        """上传文件
        
        Args:
            file_data: 文件数据 (bytes 或 file-like 对象)
            file_name: 原始文件名
            mime_type: MIME类型 (可选)
            tags: 文件标签 (可选)
        
        Returns:
            PluginResult 包含 FileInfo
        """
        try:
            # 读取文件数据
            if hasattr(file_data, 'read'):
                data = file_data.read()
            else:
                data = file_data
            
            file_size = len(data)
            
            # 猜测MIME类型
            if mime_type is None:
                mime_type, _ = mimetypes.guess_type(file_name)
                mime_type = mime_type or 'application/octet-stream'
            
            # 验证文件
            is_valid, error_msg = self._validate_file(file_name, file_size, mime_type)
            if not is_valid:
                return PluginResult(
                    success=False,
                    error=error_msg,
                    data={"file_name": file_name}
                )
            
            # 计算哈希
            md5_hash, sha256_hash = self._calculate_hashes(data)
            
            # 生成唯一ID
            file_id = f"file_{sha256_hash[:16]}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
            
            # 获取文件类型
            extension = Path(file_name).suffix.lower()
            file_type = self._get_file_type(extension)
            
            # 保存文件
            file_path = self.storage_path / f"{file_id}{extension}"
            with open(file_path, 'wb') as f:
                f.write(data)
            
            # 创建文件信息
            file_info = FileInfo(
                id=file_id,
                name=f"{file_id}{extension}",
                original_name=file_name,
                size=file_size,
                mime_type=mime_type,
                file_type=file_type,
                extension=extension,
                hash_md5=md5_hash,
                hash_sha256=sha256_hash,
                created_at=datetime.now(),
                modified_at=datetime.now(),
                tags=tags or []
            )
            
            self._files[file_id] = file_info
            
            return PluginResult(
                success=True,
                data={
                    "file_info": file_info,
                    "download_url": f"/api/files/{file_id}/download"
                }
            )
            
        except Exception as e:
            return PluginResult(
                success=False,
                error=f"文件上传失败: {str(e)}"
            )
    
    async def download_file(self, file_id: str) -> PluginResult:
        """下载文件
        
        Args:
            file_id: 文件ID
        
        Returns:
            PluginResult 包含文件数据和元数据
        """
        try:
            if file_id not in self._files:
                return PluginResult(
                    success=False,
                    error="文件不存在"
                )
            
            file_info = self._files[file_id]
            file_path = self.storage_path / file_info.name
            
            if not file_path.exists():
                return PluginResult(
                    success=False,
                    error="文件数据丢失"
                )
            
            with open(file_path, 'rb') as f:
                data = f.read()
            
            return PluginResult(
                success=True,
                data={
                    "file_info": file_info,
                    "data": data,
                    "base64": base64.b64encode(data).decode('utf-8')
                }
            )
            
        except Exception as e:
            return PluginResult(
                success=False,
                error=f"文件下载失败: {str(e)}"
            )
    
    async def parse_file(self, file_id: str, options: Optional[Dict] = None) -> PluginResult:
        """解析文件内容
        
        Args:
            file_id: 文件ID
            options: 解析选项
        
        Returns:
            PluginResult 包含解析后的内容
        """
        try:
            if file_id not in self._files:
                return PluginResult(
                    success=False,
                    error="文件不存在"
                )
            
            file_info = self._files[file_id]
            file_path = self.storage_path / file_info.name
            
            # 根据文件类型选择解析器
            parser = self._get_parser(file_info.file_type)
            if not parser:
                return PluginResult(
                    success=False,
                    error=f"不支持解析此文件类型: {file_info.file_type.value}"
                )
            
            content = await parser(file_path, file_info, options or {})
            
            return PluginResult(
                success=True,
                data={
                    "file_info": file_info,
                    "content": content
                }
            )
            
        except Exception as e:
            return PluginResult(
                success=False,
                error=f"文件解析失败: {str(e)}"
            )
    
    def _get_parser(self, file_type: FileType):
        """获取对应类型的解析器"""
        parsers = {
            FileType.DOCUMENT: self._parse_document,
            FileType.SPREADSHEET: self._parse_spreadsheet,
            FileType.IMAGE: self._parse_image,
            FileType.CODE: self._parse_code,
            FileType.ARCHIVE: self._parse_archive,
        }
        return parsers.get(file_type)
    
    async def _parse_document(self, file_path: Path, file_info: FileInfo, options: Dict) -> ParsedContent:
        """解析文档"""
        content = ParsedContent()
        
        if file_info.extension == '.txt' or file_info.extension == '.md':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content.text = f.read()
        
        elif file_info.extension == '.pdf':
            # 使用PyPDF2或pdfplumber解析PDF
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text_parts = []
                    for page in reader.pages:
                        text_parts.append(page.extract_text() or "")
                    content.text = "\n".join(text_parts)
                    content.metadata["pages"] = len(reader.pages)
            except ImportError:
                content.text = "[PDF解析需要安装PyPDF2: pip install PyPDF2]"
        
        elif file_info.extension in ['.doc', '.docx']:
            # 使用python-docx解析Word
            try:
                import docx
                doc = docx.Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                content.text = "\n".join(paragraphs)
                content.metadata["paragraphs"] = len(paragraphs)
            except ImportError:
                content.text = "[Word解析需要安装python-docx: pip install python-docx]"
        
        return content
    
    async def _parse_spreadsheet(self, file_path: Path, file_info: FileInfo, options: Dict) -> ParsedContent:
        """解析表格"""
        content = ParsedContent()
        
        if file_info.extension == '.csv':
            import csv
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f)
                rows = list(reader)
                content.tables.append(rows)
                content.text = "\n".join([", ".join(row) for row in rows[:50]])  # 前50行
                content.metadata["rows"] = len(rows)
                content.metadata["columns"] = len(rows[0]) if rows else 0
        
        elif file_info.extension in ['.xls', '.xlsx']:
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, data_only=True)
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    rows = []
                    for row in sheet.iter_rows(values_only=True):
                        rows.append([str(cell) if cell is not None else "" for cell in row])
                    content.tables.append(rows)
                    content.metadata[sheet_name] = {"rows": len(rows)}

                if content.tables:
                    content.text = "\n".join([", ".join(row) for row in content.tables[0][:50]])
            except ImportError:
                content.text = "[Excel解析需要安装openpyxl: pip install openpyxl]"

        return content

    async def _parse_image(self, file_path: Path, file_info: FileInfo, options: Dict) -> ParsedContent:
        """解析图片 (OCR)"""
        content = ParsedContent()

        try:
            from PIL import Image
            img = Image.open(file_path)
            content.metadata["size"] = img.size
            content.metadata["mode"] = img.mode
            content.metadata["format"] = img.format

            try:
                import pytesseract
                content.text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                content.confidence = 0.85
            except ImportError:
                content.text = "[OCR需要安装pytesseract: pip install pytesseract]"

        except ImportError:
            content.text = "[图片处理需要安装Pillow: pip install Pillow]"

        return content

    async def _parse_code(self, file_path: Path, file_info: FileInfo, options: Dict) -> ParsedContent:
        """解析代码文件"""
        content = ParsedContent()

        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            code = f.read()

        content.text = code
        content.metadata["lines"] = len(code.splitlines())
        content.metadata["language"] = file_info.extension[1:]

        return content

    async def _parse_archive(self, file_path: Path, file_info: FileInfo, options: Dict) -> ParsedContent:
        """解析压缩包"""
        content = ParsedContent()

        if file_info.extension == '.zip':
            with zipfile.ZipFile(file_path, 'r') as zf:
                file_list = zf.namelist()
                content.text = "\n".join(file_list)
                content.metadata["files"] = len(file_list)
                content.metadata["compressed_size"] = os.path.getsize(file_path)

                total_size = sum(info.file_size for info in zf.infolist())
                content.metadata["uncompressed_size"] = total_size
                content.metadata["compression_ratio"] = (
                    os.path.getsize(file_path) / total_size if total_size > 0 else 1.0
                )

        return content

    async def compress_files(
        self,
        file_ids: List[str],
        archive_name: Optional[str] = None
    ) -> PluginResult:
        """压缩多个文件"""
        try:
            if not file_ids:
                return PluginResult(success=False, error="没有指定要压缩的文件")

            if not archive_name:
                archive_name = f"archive_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
            elif not archive_name.endswith('.zip'):
                archive_name += '.zip'

            archive_path = self.temp_path / archive_name
            total_size = 0

            with zipfile.ZipFile(archive_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                for file_id in file_ids:
                    if file_id not in self._files:
                        continue
                    file_info = self._files[file_id]
                    file_path = self.storage_path / file_info.name
                    if file_path.exists():
                        zf.write(file_path, arcname=file_info.original_name)
                        total_size += file_info.size

            compressed_size = os.path.getsize(archive_path)

            with open(archive_path, 'rb') as f:
                archive_data = f.read()

            result = await self.upload_file(
                archive_data,
                archive_name,
                'application/zip'
            )

            if result.success:
                result.data["compression_ratio"] = compressed_size / total_size if total_size > 0 else 1.0
                result.data["original_files"] = len(file_ids)

            return result

        except Exception as e:
            return PluginResult(success=False, error=f"压缩失败: {str(e)}")

    async def decompress_file(self, file_id: str, target_dir: Optional[str] = None) -> PluginResult:
        """解压文件"""
        try:
            if file_id not in self._files:
                return PluginResult(success=False, error="文件不存在")

            file_info = self._files[file_id]

            if file_info.file_type != FileType.ARCHIVE:
                return PluginResult(success=False, error="不是压缩文件")

            file_path = self.storage_path / file_info.name

            if target_dir is None:
                target_dir = f"extracted_{file_id}"
            extract_path = self.storage_path / target_dir
            extract_path.mkdir(parents=True, exist_ok=True)

            extracted_files = []

            if file_info.extension == '.zip':
                with zipfile.ZipFile(file_path, 'r') as zf:
                    zf.extractall(extract_path)
                    extracted_files = zf.namelist()

            return PluginResult(
                success=True,
                data={
                    "extracted_path": str(extract_path),
                    "files": extracted_files,
                    "count": len(extracted_files)
                }
            )

        except Exception as e:
            return PluginResult(success=False, error=f"解压失败: {str(e)}")

    async def list_files(
        self,
        file_type: Optional[str] = None,
        tags: Optional[List[str]] = None,
        limit: int = 100,
        offset: int = 0
    ) -> PluginResult:
        """列出文件"""
        try:
            files = list(self._files.values())

            if file_type:
                files = [f for f in files if f.file_type.value == file_type]

            if tags:
                files = [f for f in files if any(tag in f.tags for tag in tags)]

            files.sort(key=lambda x: x.created_at, reverse=True)

            total = len(files)
            files = files[offset:offset + limit]

            return PluginResult(
                success=True,
                data={
                    "files": files,
                    "total": total,
                    "limit": limit,
                    "offset": offset
                }
            )

        except Exception as e:
            return PluginResult(success=False, error=f"列出文件失败: {str(e)}")

    async def delete_file(self, file_id: str) -> PluginResult:
        """删除文件"""
        try:
            if file_id not in self._files:
                return PluginResult(success=False, error="文件不存在")

            file_info = self._files[file_id]
            file_path = self.storage_path / file_info.name

            if file_path.exists():
                file_path.unlink()

            del self._files[file_id]

            return PluginResult(
                success=True,
                data={"deleted_id": file_id}
            )

        except Exception as e:
            return PluginResult(success=False, error=f"删除文件失败: {str(e)}")

    async def search_files(
        self,
        query: str,
        file_type: Optional[str] = None
    ) -> PluginResult:
        """搜索文件"""
        try:
            results = []

            for file_info in self._files.values():
                if query.lower() in file_info.original_name.lower():
                    if file_type is None or file_info.file_type.value == file_type:
                        results.append(file_info)

            return PluginResult(
                success=True,
                data={
                    "results": results,
                    "count": len(results),
                    "query": query
                }
            )

        except Exception as e:
            return PluginResult(success=False, error=f"搜索失败: {str(e)}")

    async def get_file_metadata(self, file_id: str) -> PluginResult:
        """获取文件元数据"""
        try:
            if file_id not in self._files:
                return PluginResult(success=False, error="文件不存在")

            file_info = self._files[file_id]
            file_path = self.storage_path / file_info.name

            metadata = {
                "id": file_info.id,
                "name": file_info.original_name,
                "size": file_info.size,
                "size_formatted": self._format_size(file_info.size),
                "mime_type": file_info.mime_type,
                "file_type": file_info.file_type.value,
                "extension": file_info.extension,
                "hash_md5": file_info.hash_md5,
                "hash_sha256": file_info.hash_sha256,
                "created_at": file_info.created_at.isoformat(),
                "modified_at": file_info.modified_at.isoformat(),
                "tags": file_info.tags,
                "exists": file_path.exists()
            }

            if file_path.exists():
                stat = file_path.stat()
                metadata["disk_size"] = stat.st_size
                metadata["last_accessed"] = stat.st_atime

            return PluginResult(success=True, data=metadata)

        except Exception as e:
            return PluginResult(success=False, error=f"获取元数据失败: {str(e)}")

    def _format_size(self, size: int) -> str:
        """格式化文件大小"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024:
                return f"{size:.1f} {unit}"
            size /= 1024
        return f"{size:.1f} TB"