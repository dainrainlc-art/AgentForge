"""
AgentForge Personal Brand Module
个人品牌同步管理 - 简历自动更新
"""
from typing import Dict, Any, List, Optional
from datetime import datetime
from enum import Enum
from pydantic import BaseModel, Field
from pathlib import Path
import json
import yaml
from loguru import logger

from agentforge.config import settings


class ResumeSection(str, Enum):
    SUMMARY = "summary"
    EXPERIENCE = "experience"
    EDUCATION = "education"
    SKILLS = "skills"
    PROJECTS = "projects"
    CERTIFICATIONS = "certifications"
    LANGUAGES = "languages"
    AWARDS = "awards"


class ResumeFormat(str, Enum):
    JSON = "json"
    YAML = "yaml"
    MARKDOWN = "markdown"
    HTML = "html"
    PDF = "pdf"


class Experience(BaseModel):
    id: str
    company: str
    position: str
    location: Optional[str] = None
    start_date: str
    end_date: Optional[str] = None
    current: bool = False
    description: List[str] = Field(default_factory=list)
    technologies: List[str] = Field(default_factory=list)
    achievements: List[str] = Field(default_factory=list)


class Education(BaseModel):
    id: str
    institution: str
    degree: str
    field: str
    location: Optional[str] = None
    start_date: str
    end_date: Optional[str] = None
    gpa: Optional[float] = None
    honors: List[str] = Field(default_factory=list)


class Project(BaseModel):
    id: str
    name: str
    description: str
    url: Optional[str] = None
    technologies: List[str] = Field(default_factory=list)
    highlights: List[str] = Field(default_factory=list)
    start_date: Optional[str] = None
    end_date: Optional[str] = None


class Skill(BaseModel):
    name: str
    level: str = "intermediate"
    category: str = "technical"
    years: Optional[int] = None


class Certification(BaseModel):
    id: str
    name: str
    issuer: str
    date: str
    expiry: Optional[str] = None
    credential_id: Optional[str] = None
    url: Optional[str] = None


class Resume(BaseModel):
    id: str
    name: str
    email: str
    phone: Optional[str] = None
    location: Optional[str] = None
    website: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    
    summary: str = ""
    experiences: List[Experience] = Field(default_factory=list)
    education: List[Education] = Field(default_factory=list)
    skills: List[Skill] = Field(default_factory=list)
    projects: List[Project] = Field(default_factory=list)
    certifications: List[Certification] = Field(default_factory=list)
    languages: List[Dict[str, str]] = Field(default_factory=list)
    awards: List[str] = Field(default_factory=list)
    
    created_at: datetime = Field(default_factory=datetime.now)
    updated_at: datetime = Field(default_factory=datetime.now)
    version: int = 1


class ResumeManager:
    def __init__(
        self,
        storage_path: Optional[str] = None
    ):
        self.storage_path = Path(storage_path or getattr(settings, 'resume_storage_path', 'data/resumes'))
        self.storage_path.mkdir(parents=True, exist_ok=True)
        
        self._resumes: Dict[str, Resume] = {}
        self._load_resumes()
    
    def _load_resumes(self):
        for resume_file in self.storage_path.glob("*.json"):
            try:
                with open(resume_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    resume = Resume(**data)
                    self._resumes[resume.id] = resume
            except Exception as e:
                logger.error(f"Failed to load resume {resume_file}: {e}")
    
    def _save_resume(self, resume: Resume) -> bool:
        try:
            resume_file = self.storage_path / f"{resume.id}.json"
            with open(resume_file, 'w', encoding='utf-8') as f:
                json.dump(resume.model_dump(), f, indent=2, default=str)
            return True
        except Exception as e:
            logger.error(f"Failed to save resume: {e}")
            return False
    
    def create_resume(
        self,
        name: str,
        email: str,
        **kwargs
    ) -> Resume:
        import secrets
        resume_id = secrets.token_hex(4)
        
        resume = Resume(
            id=resume_id,
            name=name,
            email=email,
            **kwargs
        )
        
        self._resumes[resume_id] = resume
        self._save_resume(resume)
        
        logger.info(f"Created resume: {resume_id}")
        return resume
    
    def get_resume(self, resume_id: str) -> Optional[Resume]:
        return self._resumes.get(resume_id)
    
    def list_resumes(self) -> List[Resume]:
        return list(self._resumes.values())
    
    def update_resume(
        self,
        resume_id: str,
        **updates
    ) -> Optional[Resume]:
        resume = self._resumes.get(resume_id)
        if not resume:
            return None
        
        for key, value in updates.items():
            if hasattr(resume, key):
                setattr(resume, key, value)
        
        resume.updated_at = datetime.now()
        resume.version += 1
        
        self._save_resume(resume)
        return resume
    
    def add_experience(
        self,
        resume_id: str,
        experience: Experience
    ) -> Optional[Resume]:
        resume = self._resumes.get(resume_id)
        if not resume:
            return None
        
        resume.experiences.append(experience)
        resume.updated_at = datetime.now()
        resume.version += 1
        
        self._save_resume(resume)
        return resume
    
    def update_experience(
        self,
        resume_id: str,
        experience_id: str,
        **updates
    ) -> Optional[Resume]:
        resume = self._resumes.get(resume_id)
        if not resume:
            return None
        
        for i, exp in enumerate(resume.experiences):
            if exp.id == experience_id:
                for key, value in updates.items():
                    if hasattr(exp, key):
                        setattr(exp, key, value)
                break
        
        resume.updated_at = datetime.now()
        resume.version += 1
        
        self._save_resume(resume)
        return resume
    
    def add_project(
        self,
        resume_id: str,
        project: Project
    ) -> Optional[Resume]:
        resume = self._resumes.get(resume_id)
        if not resume:
            return None
        
        resume.projects.append(project)
        resume.updated_at = datetime.now()
        resume.version += 1
        
        self._save_resume(resume)
        return resume
    
    def update_skills(
        self,
        resume_id: str,
        skills: List[Skill]
    ) -> Optional[Resume]:
        resume = self._resumes.get(resume_id)
        if not resume:
            return None
        
        resume.skills = skills
        resume.updated_at = datetime.now()
        resume.version += 1
        
        self._save_resume(resume)
        return resume
    
    def export_resume(
        self,
        resume_id: str,
        format: ResumeFormat = ResumeFormat.JSON
    ) -> Optional[str]:
        resume = self._resumes.get(resume_id)
        if not resume:
            return None
        
        if format == ResumeFormat.JSON:
            return json.dumps(resume.model_dump(), indent=2, default=str)
        
        elif format == ResumeFormat.YAML:
            return yaml.dump(resume.model_dump(), allow_unicode=True, default_flow_style=False)
        
        elif format == ResumeFormat.MARKDOWN:
            return self._export_markdown(resume)
        
        elif format == ResumeFormat.HTML:
            return self._export_html(resume)
        
        return None
    
    def _export_markdown(self, resume: Resume) -> str:
        lines = [
            f"# {resume.name}",
            "",
            f"📧 {resume.email}",
        ]
        
        if resume.phone:
            lines.append(f"📞 {resume.phone}")
        if resume.location:
            lines.append(f"📍 {resume.location}")
        if resume.website:
            lines.append(f"🌐 [{resume.website}]({resume.website})")
        if resume.linkedin:
            lines.append(f"💼 [LinkedIn]({resume.linkedin})")
        if resume.github:
            lines.append(f"🐙 [GitHub]({resume.github})")
        
        lines.extend(["", "---", ""])
        
        if resume.summary:
            lines.extend([
                "## 个人简介",
                "",
                resume.summary,
                ""
            ])
        
        if resume.experiences:
            lines.extend(["## 工作经历", ""])
            for exp in resume.experiences:
                lines.extend([
                    f"### {exp.position} @ {exp.company}",
                    "",
                    f"**{exp.start_date}** - **{exp.end_date or '至今'}**",
                    ""
                ])
                for desc in exp.description:
                    lines.append(f"- {desc}")
                if exp.technologies:
                    lines.append(f"\n**技术栈**: {', '.join(exp.technologies)}")
                lines.append("")
        
        if resume.projects:
            lines.extend(["## 项目经验", ""])
            for proj in resume.projects:
                lines.extend([
                    f"### {proj.name}",
                    "",
                    proj.description,
                    ""
                ])
                if proj.technologies:
                    lines.append(f"**技术**: {', '.join(proj.technologies)}")
                lines.append("")
        
        if resume.skills:
            lines.extend(["## 技能", ""])
            skill_groups = {}
            for skill in resume.skills:
                if skill.category not in skill_groups:
                    skill_groups[skill.category] = []
                skill_groups[skill.category].append(skill.name)
            
            for category, skills in skill_groups.items():
                lines.append(f"- **{category.title()}**: {', '.join(skills)}")
            lines.append("")
        
        if resume.education:
            lines.extend(["## 教育背景", ""])
            for edu in resume.education:
                lines.append(f"- **{edu.degree}** in {edu.field} - {edu.institution} ({edu.start_date} - {edu.end_date or '至今'})")
            lines.append("")
        
        if resume.certifications:
            lines.extend(["## 认证证书", ""])
            for cert in resume.certifications:
                lines.append(f"- **{cert.name}** - {cert.issuer} ({cert.date})")
            lines.append("")
        
        return "\n".join(lines)
    
    def _export_html(self, resume: Resume) -> str:
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{resume.name} - 简历</title>
    <style>
        body {{ font-family: 'Segoe UI', sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1 {{ color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; margin-top: 30px; }}
        h3 {{ color: #7f8c8d; }}
        .contact {{ color: #7f8c8d; margin-bottom: 20px; }}
        .experience {{ margin-bottom: 20px; }}
        .skills {{ display: flex; flex-wrap: wrap; gap: 10px; }}
        .skill {{ background: #ecf0f1; padding: 5px 10px; border-radius: 5px; }}
    </style>
</head>
<body>
    <h1>{resume.name}</h1>
    <div class="contact">
        {resume.email}
        {f' | {resume.phone}' if resume.phone else ''}
        {f' | {resume.location}' if resume.location else ''}
    </div>
    
    <h2>个人简介</h2>
    <p>{resume.summary}</p>
"""
        
        if resume.experiences:
            html += "<h2>工作经历</h2>"
            for exp in resume.experiences:
                html += f"""
    <div class="experience">
        <h3>{exp.position} @ {exp.company}</h3>
        <p>{exp.start_date} - {exp.end_date or '至今'}</p>
        <ul>
            {''.join(f'<li>{d}</li>' for d in exp.description)}
        </ul>
    </div>
"""
        
        if resume.skills:
            html += "<h2>技能</h2><div class='skills'>"
            for skill in resume.skills:
                html += f"<span class='skill'>{skill.name}</span>"
            html += "</div>"
        
        html += "</body></html>"
        return html


class SkillSuggestion(BaseModel):
    skill_name: str
    category: str
    reason: str
    relevance_score: float = 0.0
    source: str = "analysis"
    trending: bool = False


class LinkedInSync:
    def __init__(self):
        self.resume_manager = ResumeManager()
        self._skill_database = self._load_skill_database()
    
    def _load_skill_database(self) -> Dict[str, Dict[str, Any]]:
        return {
            "python": {"category": "technical", "trending": True, "related": ["django", "flask", "fastapi", "pandas", "numpy"]},
            "javascript": {"category": "technical", "trending": True, "related": ["react", "vue", "angular", "node.js", "typescript"]},
            "typescript": {"category": "technical", "trending": True, "related": ["javascript", "react", "angular", "node.js"]},
            "react": {"category": "technical", "trending": True, "related": ["javascript", "typescript", "next.js", "redux"]},
            "node.js": {"category": "technical", "trending": True, "related": ["javascript", "typescript", "express", "nestjs"]},
            "docker": {"category": "devops", "trending": True, "related": ["kubernetes", "ci/cd", "containerization"]},
            "kubernetes": {"category": "devops", "trending": True, "related": ["docker", "helm", "container orchestration"]},
            "aws": {"category": "cloud", "trending": True, "related": ["ec2", "s3", "lambda", "cloudformation"]},
            "machine learning": {"category": "ai", "trending": True, "related": ["tensorflow", "pytorch", "scikit-learn", "deep learning"]},
            "generative ai": {"category": "ai", "trending": True, "related": ["llm", "chatgpt", "stable diffusion", "prompt engineering"]},
            "sql": {"category": "technical", "trending": False, "related": ["postgresql", "mysql", "database design"]},
            "git": {"category": "technical", "trending": False, "related": ["github", "gitlab", "version control"]},
            "api design": {"category": "technical", "trending": True, "related": ["rest", "graphql", "openapi"]},
            "agile": {"category": "methodology", "trending": False, "related": ["scrum", "kanban", "sprint planning"]},
            "project management": {"category": "soft", "trending": False, "related": ["leadership", "communication", "team coordination"]},
        }
    
    async def sync_profile_to_resume(
        self,
        resume_id: str,
        linkedin_data: Dict[str, Any]
    ) -> Optional[Resume]:
        resume = self.resume_manager.get_resume(resume_id)
        if not resume:
            return None
        
        if "summary" in linkedin_data:
            resume.summary = linkedin_data["summary"]
        
        if "experiences" in linkedin_data:
            for exp_data in linkedin_data["experiences"]:
                experience = Experience(
                    id=exp_data.get("id", ""),
                    company=exp_data.get("company", ""),
                    position=exp_data.get("title", ""),
                    location=exp_data.get("location"),
                    start_date=exp_data.get("startDate", ""),
                    end_date=exp_data.get("endDate"),
                    current=exp_data.get("current", False),
                    description=exp_data.get("description", "").split("\n") if exp_data.get("description") else []
                )
                resume.experiences.append(experience)
        
        if "skills" in linkedin_data:
            skills = [
                Skill(name=skill, level="intermediate", category="technical")
                for skill in linkedin_data["skills"]
            ]
            resume.skills = skills
        
        resume.updated_at = datetime.now()
        resume.version += 1
        
        self.resume_manager._save_resume(resume)
        return resume
    
    async def suggest_skill_updates(
        self,
        resume_id: str,
        recent_work: List[Dict[str, Any]]
    ) -> List[str]:
        resume = self.resume_manager.get_resume(resume_id)
        if not resume:
            return []
        
        existing_skills = {s.name.lower() for s in resume.skills}
        suggested_skills = []
        
        for work in recent_work:
            if "technologies" in work:
                for tech in work["technologies"]:
                    if tech.lower() not in existing_skills:
                        suggested_skills.append(tech)
        
        return list(set(suggested_skills))
    
    async def analyze_and_suggest_skills(
        self,
        resume_id: str,
        industry: Optional[str] = None,
        target_role: Optional[str] = None
    ) -> List[SkillSuggestion]:
        from agentforge.llm.model_router import ModelRouter
        
        resume = self.resume_manager.get_resume(resume_id)
        if not resume:
            return []
        
        existing_skills = {s.name.lower() for s in resume.skills}
        suggestions = []
        
        for skill_name, skill_info in self._skill_database.items():
            if skill_name in existing_skills:
                continue
            
            related_existing = any(
                related.lower() in existing_skills
                for related in skill_info.get("related", [])
            )
            
            if related_existing:
                relevance_score = 0.8 if skill_info.get("trending") else 0.6
                suggestions.append(SkillSuggestion(
                    skill_name=skill_name.title(),
                    category=skill_info["category"],
                    reason=f"Related to your existing skills: {', '.join(skill_info['related'][:3])}",
                    relevance_score=relevance_score,
                    source="related_skills",
                    trending=skill_info.get("trending", False)
                ))
        
        if resume.experiences:
            for exp in resume.experiences[:3]:
                for tech in exp.technologies:
                    tech_lower = tech.lower()
                    if tech_lower not in existing_skills:
                        if tech_lower in self._skill_database:
                            skill_info = self._skill_database[tech_lower]
                            suggestions.append(SkillSuggestion(
                                skill_name=tech,
                                category=skill_info["category"],
                                reason=f"Used in your role at {exp.company}",
                                relevance_score=0.9,
                                source="experience",
                                trending=skill_info.get("trending", False)
                            ))
                        else:
                            suggestions.append(SkillSuggestion(
                                skill_name=tech,
                                category="technical",
                                reason=f"Used in your role at {exp.company}",
                                relevance_score=0.7,
                                source="experience",
                                trending=False
                            ))
        
        if target_role or industry:
            llm = ModelRouter()
            prompt = f"""Based on this resume, suggest additional skills that would be valuable.

Current Skills: {', '.join(s.name for s in resume.skills)}
Industry: {industry or 'Technology'}
Target Role: {target_role or 'Not specified'}

Recent Experience:
{chr(10).join(f"- {e.position} at {e.company}" for e in resume.experiences[:3])}

Suggest 5-10 skills that would enhance this profile. Format as JSON array:
[{{"skill": "name", "category": "category", "reason": "why it's relevant"}}]"""

            try:
                response = await llm.chat_with_failover(
                    message=prompt,
                    task_type="analysis"
                )
                
                import json
                ai_suggestions = json.loads(response)
                
                for sug in ai_suggestions:
                    skill_name = sug.get("skill", "")
                    if skill_name.lower() not in existing_skills:
                        suggestions.append(SkillSuggestion(
                            skill_name=skill_name,
                            category=sug.get("category", "technical"),
                            reason=sug.get("reason", "AI suggested"),
                            relevance_score=0.75,
                            source="ai_analysis",
                            trending=False
                        ))
            except Exception as e:
                logger.warning(f"AI skill suggestion failed: {e}")
        
        seen = set()
        unique_suggestions = []
        for sug in suggestions:
            key = sug.skill_name.lower()
            if key not in seen:
                seen.add(key)
                unique_suggestions.append(sug)
        
        unique_suggestions.sort(key=lambda x: x.relevance_score, reverse=True)
        
        return unique_suggestions[:15]
    
    async def get_trending_skills(self, category: Optional[str] = None) -> List[SkillSuggestion]:
        trending = []
        for skill_name, skill_info in self._skill_database.items():
            if skill_info.get("trending"):
                if category is None or skill_info["category"] == category:
                    trending.append(SkillSuggestion(
                        skill_name=skill_name.title(),
                        category=skill_info["category"],
                        reason="Currently in high demand",
                        relevance_score=0.9,
                        source="trending",
                        trending=True
                    ))
        
        return trending
    
    async def get_skill_gap_analysis(
        self,
        resume_id: str,
        target_role: str
    ) -> Dict[str, Any]:
        from agentforge.llm.model_router import ModelRouter
        
        resume = self.resume_manager.get_resume(resume_id)
        if not resume:
            return {"error": "Resume not found"}
        
        llm = ModelRouter()
        
        prompt = f"""Analyze the skill gap for this candidate targeting a {target_role} role.

Current Skills:
{chr(10).join(f'- {s.name} ({s.level}, {s.category})' for s in resume.skills)}

Experience Summary:
{chr(10).join(f'- {e.position} at {e.company}: {", ".join(e.technologies[:5])}' for e in resume.experiences[:3])}

Provide a skill gap analysis in JSON format:
{{
    "matching_skills": ["skills that match the target role"],
    "missing_critical": ["critical skills missing"],
    "missing_nice_to_have": ["helpful but not critical skills"],
    "recommended_learning_path": ["ordered list of skills to learn"],
    "estimated_effort": "low/medium/high"
}}"""

        try:
            response = await llm.chat_with_failover(
                message=prompt,
                task_type="analysis"
            )
            
            import json
            return json.loads(response)
        except Exception as e:
            logger.error(f"Skill gap analysis failed: {e}")
            return {"error": str(e)}


class FiverrProfileSync:
    def __init__(self):
        self.resume_manager = ResumeManager()
    
    async def sync_gig_to_projects(
        self,
        resume_id: str,
        gig_data: Dict[str, Any]
    ) -> Optional[Resume]:
        resume = self.resume_manager.get_resume(resume_id)
        if not resume:
            return None
        
        project = Project(
            id=gig_data.get("id", ""),
            name=gig_data.get("title", ""),
            description=gig_data.get("description", ""),
            url=gig_data.get("url"),
            technologies=gig_data.get("tags", []),
            highlights=gig_data.get("highlights", [])
        )
        
        resume.projects.append(project)
        resume.updated_at = datetime.now()
        resume.version += 1
        
        self.resume_manager._save_resume(resume)
        return resume
    
    async def generate_gig_description_from_resume(
        self,
        resume_id: str,
        service_type: str
    ) -> str:
        resume = self.resume_manager.get_resume(resume_id)
        if not resume:
            return ""
        
        relevant_skills = [
            s.name for s in resume.skills
            if s.category in ["technical", service_type]
        ]
        
        recent_experience = resume.experiences[0] if resume.experiences else None
        
        description = f"Experienced {service_type} professional"
        
        if relevant_skills:
            description += f" with expertise in {', '.join(relevant_skills[:5])}"
        
        if recent_experience:
            description += f". Currently working as {recent_experience.position} at {recent_experience.company}"
        
        return description


class ResumeAutoUpdater:
    def __init__(self):
        self.resume_manager = ResumeManager()
        self.llm = None
    
    def _get_llm(self):
        if self.llm is None:
            from agentforge.llm.model_router import ModelRouter
            self.llm = ModelRouter()
        return self.llm
    
    async def auto_update_from_project(
        self,
        resume_id: str,
        project_data: Dict[str, Any]
    ) -> Optional[Resume]:
        resume = self.resume_manager.get_resume(resume_id)
        if not resume:
            return None
        
        new_skills = project_data.get("technologies", [])
        existing_skill_names = {s.name.lower() for s in resume.skills}
        
        for skill_name in new_skills:
            if skill_name.lower() not in existing_skill_names:
                resume.skills.append(Skill(
                    name=skill_name,
                    level="intermediate",
                    category="technical"
                ))
        
        project = Project(
            id=project_data.get("id", f"proj_{datetime.now().strftime('%Y%m%d%H%M%S')}"),
            name=project_data.get("name", ""),
            description=project_data.get("description", ""),
            url=project_data.get("url"),
            technologies=project_data.get("technologies", []),
            highlights=project_data.get("highlights", []),
            start_date=project_data.get("start_date"),
            end_date=project_data.get("end_date")
        )
        
        existing_project_ids = {p.id for p in resume.projects}
        if project.id not in existing_project_ids:
            resume.projects.append(project)
        
        resume.updated_at = datetime.now()
        resume.version += 1
        
        self.resume_manager._save_resume(resume)
        logger.info(f"Auto-updated resume {resume_id} from project {project.name}")
        
        return resume
    
    async def auto_update_from_experience(
        self,
        resume_id: str,
        experience_data: Dict[str, Any]
    ) -> Optional[Resume]:
        resume = self.resume_manager.get_resume(resume_id)
        if not resume:
            return None
        
        experience = Experience(
            id=experience_data.get("id", f"exp_{datetime.now().strftime('%Y%m%d%H%M%S')}"),
            company=experience_data.get("company", ""),
            position=experience_data.get("position", ""),
            location=experience_data.get("location"),
            start_date=experience_data.get("start_date", ""),
            end_date=experience_data.get("end_date"),
            current=experience_data.get("current", False),
            description=experience_data.get("description", []),
            technologies=experience_data.get("technologies", []),
            achievements=experience_data.get("achievements", [])
        )
        
        existing_exp_ids = {e.id for e in resume.experiences}
        if experience.id not in existing_exp_ids:
            for tech in experience.technologies:
                existing_skill_names = {s.name.lower() for s in resume.skills}
                if tech.lower() not in existing_skill_names:
                    resume.skills.append(Skill(
                        name=tech,
                        level="intermediate",
                        category="technical"
                    ))
            
            resume.experiences.insert(0, experience)
        
        resume.updated_at = datetime.now()
        resume.version += 1
        
        self.resume_manager._save_resume(resume)
        return resume
    
    async def generate_summary_update(
        self,
        resume_id: str
    ) -> str:
        resume = self.resume_manager.get_resume(resume_id)
        if not resume:
            return ""
        
        llm = self._get_llm()
        
        experience_summary = ""
        if resume.experiences:
            exp = resume.experiences[0]
            experience_summary = f"Currently {exp.position} at {exp.company}"
        
        skills_summary = ", ".join([s.name for s in resume.skills[:5]])
        
        prompt = f"""Generate a professional resume summary based on this information:

Current Role: {experience_summary}
Top Skills: {skills_summary}
Total Experience: {len(resume.experiences)} positions
Projects: {len(resume.projects)} projects

Create a 2-3 sentence professional summary that highlights key strengths and career focus.

Summary:"""

        response = await llm.chat_with_failover(
            message=prompt,
            task_type="creative"
        )
        
        return response.strip()
    
    async def suggest_achievements(
        self,
        resume_id: str,
        experience_id: str
    ) -> List[str]:
        resume = self.resume_manager.get_resume(resume_id)
        if not resume:
            return []
        
        experience = None
        for exp in resume.experiences:
            if exp.id == experience_id:
                experience = exp
                break
        
        if not experience:
            return []
        
        llm = self._get_llm()
        
        prompt = f"""Based on this work experience, suggest 3-5 professional achievements:

Position: {experience.position}
Company: {experience.company}
Description: {'. '.join(experience.description)}
Technologies: {', '.join(experience.technologies)}

Generate achievements in the format "Action verb + Result + Metric" (e.g., "Developed automation system that reduced processing time by 40%")

Achievements:"""

        response = await llm.chat_with_failover(
            message=prompt,
            task_type="creative"
        )
        
        achievements = []
        for line in response.strip().split('\n'):
            line = line.strip()
            if line and not line.startswith('#'):
                line = line.lstrip('0123456789.-) ')
                if line:
                    achievements.append(line)
        
        return achievements[:5]
    
    async def review_and_suggest(
        self,
        resume_id: str
    ) -> Dict[str, Any]:
        resume = self.resume_manager.get_resume(resume_id)
        if not resume:
            return {"error": "Resume not found"}
        
        llm = self._get_llm()
        
        issues = []
        suggestions = []
        
        if not resume.summary or len(resume.summary) < 50:
            issues.append("Summary is missing or too short")
            suggestions.append({
                "type": "summary",
                "suggestion": await self.generate_summary_update(resume_id)
            })
        
        if len(resume.skills) < 5:
            issues.append("Not enough skills listed")
            linkedin_sync = LinkedInSync()
            skill_suggestions = await linkedin_sync.analyze_and_suggest_skills(resume_id)
            suggestions.append({
                "type": "skills",
                "suggestion": [s.skill_name for s in skill_suggestions[:5]]
            })
        
        for exp in resume.experiences:
            if not exp.achievements:
                issues.append(f"No achievements for {exp.position} at {exp.company}")
                achievements = await self.suggest_achievements(resume_id, exp.id)
                suggestions.append({
                    "type": "achievements",
                    "experience_id": exp.id,
                    "suggestion": achievements
                })
        
        if not resume.projects and resume.experiences:
            issues.append("No projects listed")
            suggestions.append({
                "type": "projects",
                "suggestion": "Consider adding projects from your work experience"
            })
        
        days_since_update = (datetime.now() - resume.updated_at).days
        if days_since_update > 30:
            issues.append(f"Resume hasn't been updated in {days_since_update} days")
        
        return {
            "resume_id": resume_id,
            "last_updated": resume.updated_at.isoformat(),
            "version": resume.version,
            "issues": issues,
            "suggestions": suggestions,
            "score": self._calculate_resume_score(resume)
        }
    
    def _calculate_resume_score(self, resume: Resume) -> int:
        score = 0
        
        if resume.summary and len(resume.summary) >= 50:
            score += 15
        
        score += min(len(resume.experiences) * 10, 30)
        
        score += min(len(resume.skills) * 2, 20)
        
        for exp in resume.experiences:
            if exp.achievements:
                score += min(len(exp.achievements) * 2, 10)
        
        score += min(len(resume.projects) * 5, 15)
        
        if resume.certifications:
            score += min(len(resume.certifications) * 3, 10)
        
        return min(score, 100)


class ResumeTemplate(BaseModel):
    id: str
    name: str
    description: str
    style: str
    sections: List[str]
    color_scheme: Dict[str, str] = Field(default_factory=lambda: {"primary": "#2563eb", "secondary": "#64748b"})


class ResumeExporter:
    def __init__(self):
        self.resume_manager = ResumeManager()
    
    async def export_to_pdf(
        self,
        resume_id: str,
        template: Optional[ResumeTemplate] = None
    ) -> Optional[bytes]:
        resume = self.resume_manager.get_resume(resume_id)
        if not resume:
            return None
        
        html_content = self._generate_resume_html(resume, template)
        
        try:
            import weasyprint
            pdf = weasyprint.HTML(string=html_content).write_pdf()
            return pdf
        except ImportError:
            logger.warning("weasyprint not installed, returning HTML instead")
            return html_content.encode('utf-8')
    
    async def export_to_docx(
        self,
        resume_id: str
    ) -> Optional[bytes]:
        resume = self.resume_manager.get_resume(resume_id)
        if not resume:
            return None
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            title = doc.add_heading(resume.name, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.add_run(resume.email)
            if resume.phone:
                contact_para.add_run(f" | {resume.phone}")
            if resume.location:
                contact_para.add_run(f" | {resume.location}")
            
            if resume.summary:
                doc.add_heading("Summary", level=1)
                doc.add_paragraph(resume.summary)
            
            if resume.experiences:
                doc.add_heading("Experience", level=1)
                for exp in resume.experiences:
                    doc.add_heading(f"{exp.position} at {exp.company}", level=2)
                    date_para = doc.add_paragraph()
                    date_para.add_run(f"{exp.start_date} - {exp.end_date or 'Present'}")
                    
                    for desc in exp.description:
                        doc.add_paragraph(desc, style='List Bullet')
                    
                    if exp.technologies:
                        doc.add_paragraph(f"Technologies: {', '.join(exp.technologies)}")
            
            if resume.skills:
                doc.add_heading("Skills", level=1)
                skill_groups = {}
                for skill in resume.skills:
                    if skill.category not in skill_groups:
                        skill_groups[skill.category] = []
                    skill_groups[skill.category].append(skill.name)
                
                for category, skills in skill_groups.items():
                    doc.add_paragraph(f"{category.title()}: {', '.join(skills)}")
            
            if resume.projects:
                doc.add_heading("Projects", level=1)
                for proj in resume.projects:
                    doc.add_heading(proj.name, level=2)
                    doc.add_paragraph(proj.description)
                    if proj.technologies:
                        doc.add_paragraph(f"Technologies: {', '.join(proj.technologies)}")
            
            if resume.education:
                doc.add_heading("Education", level=1)
                for edu in resume.education:
                    doc.add_paragraph(f"{edu.degree} in {edu.field} - {edu.institution}")
            
            import io
            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
            
        except ImportError:
            logger.warning("python-docx not installed")
            return None
    
    def _generate_resume_html(
        self,
        resume: Resume,
        template: Optional[ResumeTemplate] = None
    ) -> str:
        colors = template.color_scheme if template else {"primary": "#2563eb", "secondary": "#64748b"}
        
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{resume.name} - Resume</title>
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 40px; line-height: 1.6; }}
        h1 {{ color: {colors['primary']}; border-bottom: 2px solid {colors['primary']}; padding-bottom: 10px; }}
        h2 {{ color: {colors['secondary']}; margin-top: 30px; }}
        h3 {{ color: #333; margin-bottom: 5px; }}
        .contact {{ text-align: center; color: #666; margin-bottom: 30px; }}
        .experience {{ margin-bottom: 20px; }}
        .skills {{ display: flex; flex-wrap: wrap; gap: 10px; }}
        .skill {{ background: #f1f5f9; padding: 5px 15px; border-radius: 20px; font-size: 14px; }}
        .date {{ color: #666; font-size: 14px; }}
        ul {{ margin-top: 5px; }}
        li {{ margin-bottom: 5px; }}
    </style>
</head>
<body>
    <h1>{resume.name}</h1>
    <div class="contact">
        {resume.email}
        {f' | {resume.phone}' if resume.phone else ''}
        {f' | {resume.location}' if resume.location else ''}
    </div>
"""
        
        if resume.summary:
            html += f"""
    <h2>Professional Summary</h2>
    <p>{resume.summary}</p>
"""
        
        if resume.experiences:
            html += "<h2>Experience</h2>"
            for exp in resume.experiences:
                html += f"""
    <div class="experience">
        <h3>{exp.position}</h3>
        <div class="date">{exp.company} | {exp.start_date} - {exp.end_date or 'Present'}</div>
        <ul>
            {''.join(f'<li>{d}</li>' for d in exp.description)}
        </ul>
        {f'<p><strong>Technologies:</strong> {", ".join(exp.technologies)}</p>' if exp.technologies else ''}
    </div>
"""
        
        if resume.skills:
            html += """
    <h2>Skills</h2>
    <div class="skills">
"""
            for skill in resume.skills:
                html += f'<span class="skill">{skill.name}</span>'
            html += "    </div>"
        
        if resume.projects:
            html += "<h2>Projects</h2>"
            for proj in resume.projects:
                html += f"""
    <div class="experience">
        <h3>{proj.name}</h3>
        <p>{proj.description}</p>
        {f'<p><strong>Technologies:</strong> {", ".join(proj.technologies)}</p>' if proj.technologies else ''}
    </div>
"""
        
        if resume.education:
            html += "<h2>Education</h2>"
            for edu in resume.education:
                html += f"""
    <div class="experience">
        <h3>{edu.degree} in {edu.field}</h3>
        <div class="date">{edu.institution} | {edu.start_date} - {edu.end_date or 'Present'}</div>
    </div>
"""
        
        html += "</body></html>"
        return html
